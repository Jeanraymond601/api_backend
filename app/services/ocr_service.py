import os
import tempfile
import time
import re
from typing import List, Optional, Tuple, Dict, Any
from PIL import Image
from paddleocr import PaddleOCR
import cv2
import numpy as np
import magic
import logging
from concurrent.futures import ThreadPoolExecutor, as_completed

logger = logging.getLogger(__name__)

class OCRService:
    def __init__(self, config):
        self.config = config
        self.ocr_engine = None
        self.max_workers = config.get("MAX_CONCURRENT_OCR", 3)
        self.timeout = config.get("OCR_TIMEOUT", 30)
        self._init_ocr_engine()
        
    def _init_ocr_engine(self):
        """Initialize OCR engine based on configuration"""
        engine_type = self.config.get("OCR_ENGINE", "paddleocr").lower()
        
        try:
            if engine_type == "paddleocr":
                # DÉSACTIVER LES LOGS VERBOSE POUR ACCÉLÉRER LE DÉMARRAGE
                import os
                os.environ['DISABLE_MODEL_SOURCE_CHECK'] = 'True'
                
                # CONFIGURATION ULTRA-SIMPLE POUR PADDLEOCR 3.3.2
                # Essayer différentes configurations jusqu'à ce qu'une fonctionne
                try:
                    # Essai 1: Configuration minimale (juste la langue)
                    self.ocr_engine = PaddleOCR(lang='fr')
                    logger.info("PaddleOCR initialized with minimal config (lang='fr')")
                except Exception as e1:
                    logger.warning(f"Minimal config failed: {e1}")
                    try:
                        # Essai 2: Sans aucun paramètre
                        self.ocr_engine = PaddleOCR()
                        logger.info("PaddleOCR initialized without params")
                    except Exception as e2:
                        logger.warning(f"No params config failed: {e2}")
                        # Essai 3: Ancienne configuration (pour compatibilité)
                        self.ocr_engine = PaddleOCR(
                            lang='fr',
                            use_angle_cls=False,
                            use_gpu=False
                        )
                        logger.info("PaddleOCR initialized with legacy config")
                
            elif engine_type == "tesseract":
                import pytesseract
                tesseract_path = self.config.get("TESSERACT_PATH")
                if tesseract_path and os.path.exists(tesseract_path):
                    pytesseract.pytesseract.tesseract_cmd = tesseract_path
                self.ocr_engine = pytesseract
                logger.info("Tesseract OCR initialized")
                
            elif engine_type == "easyocr":
                import easyocr
                langs = self.config.get("PADDLE_OCR_LANGS", ["fr", "en"])
                self.ocr_engine = easyocr.Reader(langs)
                logger.info(f"EasyOCR initialized with languages: {langs}")
                
            else:
                raise ValueError(f"Unsupported OCR engine: {engine_type}")
                
        except Exception as e:
            logger.error(f"Failed to initialize OCR engine: {e}")
            # Créer un mock pour éviter le crash de l'application
            class MockOCR:
                def predict(self, img):
                    return [{'rec_texts': [], 'rec_scores': []}]
                def ocr(self, img):
                    return []
            self.ocr_engine = MockOCR()
            logger.warning("Using mock OCR engine due to initialization failure")
    
    def preprocess_image(self, image_path: str) -> np.ndarray:
        """Preprocess image for better OCR results"""
        try:
            # Lire l'image avec PIL (plus fiable)
            pil_img = Image.open(image_path)
            img = np.array(pil_img)
            
            # Convertir RGB à BGR pour OpenCV si nécessaire
            if len(img.shape) == 3 and img.shape[2] == 3:
                img = cv2.cvtColor(img, cv2.COLOR_RGB2BGR)
            elif len(img.shape) == 2:  # Niveaux de gris
                img = cv2.cvtColor(img, cv2.COLOR_GRAY2BGR)
            elif len(img.shape) == 3 and img.shape[2] == 4:  # RGBA
                img = cv2.cvtColor(img, cv2.COLOR_BGRA2BGR)
            
            # Appliquer un prétraitement léger si configuré
            if self.config.get("preprocess_image", True):
                # Redimensionner si trop grande (pour accélérer l'OCR)
                max_size = 2000
                height, width = img.shape[:2]
                if max(height, width) > max_size:
                    scale = max_size / max(height, width)
                    new_width = int(width * scale)
                    new_height = int(height * scale)
                    img = cv2.resize(img, (new_width, new_height), interpolation=cv2.INTER_AREA)
                
                # Améliorer le contraste (aide l'OCR)
                lab = cv2.cvtColor(img, cv2.COLOR_BGR2LAB)
                l, a, b = cv2.split(lab)
                clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
                l = clahe.apply(l)
                lab = cv2.merge([l, a, b])
                img = cv2.cvtColor(lab, cv2.COLOR_LAB2BGR)
            
            return img
            
        except Exception as e:
            logger.error(f"Image preprocessing failed: {e}")
            # Fallback: lire directement avec OpenCV
            img = cv2.imread(image_path)
            if img is not None and len(img.shape) == 2:
                img = cv2.cvtColor(img, cv2.COLOR_GRAY2BGR)
            return img if img is not None else np.zeros((100, 100, 3), dtype=np.uint8)
    
    def extract_from_image(self, image_path: str, language: str = None) -> Tuple[str, float, float]:
        """Extract text from single image - CORRIGÉ POUR PADDLEOCR 3.3.2"""
        start_time = time.time()
        text = ""
        confidence = 0.0
        
        try:
            # Préparer l'image
            processed_img = self.preprocess_image(image_path)
            
            # OCR selon le moteur
            if isinstance(self.ocr_engine, PaddleOCR):
                # IMPORTANT: Utiliser .predict() pour PaddleOCR v3.x
                result = self.ocr_engine.predict(processed_img)
                
                if result and len(result) > 0:
                    ocr_result = result[0]  # Dictionnaire OCRResult
                    
                    # Format PaddleOCR v3.x
                    if isinstance(ocr_result, dict):
                        texts = ocr_result.get('rec_texts', [])
                        scores = ocr_result.get('rec_scores', [])
                        
                        if isinstance(texts, list) and texts:
                            # Combiner le texte
                            text_lines = []
                            confidences = []
                            
                            for i, txt in enumerate(texts):
                                txt_str = str(txt).strip()
                                if txt_str:
                                    text_lines.append(txt_str)
                                    # Ajouter la confiance correspondante
                                    if i < len(scores):
                                        try:
                                            confidences.append(float(scores[i]))
                                        except:
                                            confidences.append(0.8)  # Valeur par défaut
                            
                            if text_lines:
                                text = "\n".join(text_lines)
                                if confidences:
                                    confidence = np.mean(confidences)
                                else:
                                    confidence = 0.8  # Confiance par défaut
                                
                                logger.info(f"OCR: {len(text_lines)} lines, confidence: {confidence:.2f}")
                    else:
                        logger.warning(f"Unexpected OCR result format: {type(ocr_result)}")
            
            elif hasattr(self.ocr_engine, 'image_to_string'):  # Tesseract
                custom_config = f'--oem 3 --psm 6'
                if language:
                    custom_config += f' -l {language}'
                
                text = self.ocr_engine.image_to_string(processed_img, config=custom_config)
                confidence = 0.8
                
            elif hasattr(self.ocr_engine, 'readtext'):  # EasyOCR
                results = self.ocr_engine.readtext(processed_img)
                text = " ".join([result[1] for result in results])
                confidence = np.mean([result[2] for result in results]) if results else 0.0
            
            processing_time = time.time() - start_time
            
            if text:
                logger.debug(f"OCR completed in {processing_time:.2f}s - {len(text)} chars")
            else:
                logger.warning(f"No text extracted in {processing_time:.2f}s")
            
            return text.strip(), confidence, processing_time
            
        except Exception as e:
            logger.error(f"OCR extraction failed: {e}")
            processing_time = time.time() - start_time
            return "", 0.0, processing_time
    
    def extract_from_pdf(self, pdf_path: str, language: str = None) -> List[Tuple[str, float]]:
        """Extract text from PDF document - VERSION CORRIGÉE AVEC PyMuPDF"""
        results = []
        
        try:
            # UTILISER PyMuPDF AU LIEU DE pdf2image
            import fitz  # PyMuPDF
            from PIL import Image
            import io
            
            logger.info(f"🔍 Ouverture PDF avec PyMuPDF: {pdf_path}")
            
            # Ouvrir le PDF
            doc = fitz.open(pdf_path)
            total_pages = len(doc)
            logger.info(f"📄 PDF détecté: {total_pages} pages")
            
            # Limiter le nombre de pages pour le test
            max_pages = min(total_pages, 5) if self.config.get("PDF_MAX_PAGES", 10) else total_pages
            
            for page_num in range(max_pages):
                try:
                    logger.info(f"  Traitement page {page_num + 1}/{max_pages}...")
                    
                    page = doc[page_num]
                    
                    # Convertir la page en image avec bonne résolution
                    zoom = 2.0  # Zoom pour meilleure qualité OCR
                    mat = fitz.Matrix(zoom, zoom)
                    pix = page.get_pixmap(matrix=mat, alpha=False)
                    
                    # Convertir en image PIL
                    img_data = pix.tobytes("ppm")
                    img = Image.open(io.BytesIO(img_data))
                    
                    # Sauvegarder temporairement
                    with tempfile.NamedTemporaryFile(suffix=f'_page{page_num}.jpg', delete=False) as tmp:
                        img.save(tmp.name, 'JPEG', quality=95, optimize=True)
                        temp_path = tmp.name
                    
                    try:
                        # OCR sur l'image
                        text, confidence, _ = self.extract_from_image(temp_path, language)
                        results.append((text, confidence))
                        
                        logger.info(f"    ✅ Page {page_num + 1}: {len(text)} caractères, confiance: {confidence:.3f}")
                        
                    finally:
                        # Nettoyer le fichier temporaire
                        if os.path.exists(temp_path):
                            os.unlink(temp_path)
                            
                except Exception as e:
                    logger.error(f"Erreur page {page_num}: {e}")
                    results.append(("", 0.0))
            
            doc.close()
            logger.info(f"✅ PDF traité: {len(results)} pages extraites")
            
        except ImportError:
            logger.error("❌ PyMuPDF non installé. Exécutez: pip install PyMuPDF")
            # Fallback: essayer pdf2image si disponible (mais ça va échouer sans Poppler)
            try:
                import pdf2image
                logger.warning("⚠️  Utilisation de pdf2image (nécessite Poppler)")
                images = pdf2image.convert_from_path(pdf_path)
                logger.info(f"Fallback: PDF converti en {len(images)} pages avec pdf2image")
                
                for i, image in enumerate(images):
                    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                        image.save(tmp.name, 'PNG')
                        text, confidence, _ = self.extract_from_image(tmp.name, language)
                        results.append((text, confidence))
                        os.unlink(tmp.name)
                        
            except Exception as e:
                logger.error(f"Échec extraction PDF: {e}")
                
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
        
        return results
    
    def extract_from_docx(self, docx_path: str) -> Tuple[str, float, float]:
        """Extract text from DOCX file - VERSION AMÉLIORÉE"""
        import time
        start_time = time.time()
        
        try:
            logger.info(f"🔍 Début extraction Word: {docx_path}")
            
            # Méthode 1: python-docx (si disponible)
            text_from_docx = ""
            try:
                from docx import Document
                
                doc = Document(docx_path)
                paragraphs = []
                
                # Extraire les paragraphes
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text)
                
                # Extraire les tableaux
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text)
                        if row_text:
                            paragraphs.append(" | ".join(row_text))
                
                text_from_docx = "\n".join(paragraphs)
                logger.info(f"✅ python-docx: {len(text_from_docx)} caractères")
                
            except ImportError:
                logger.warning("⚠️  python-docx non installé")
            except Exception as e:
                logger.warning(f"⚠️  Erreur python-docx: {e}")
            
            # Méthode 2: zipfile + XML (fonctionne TOUJOURS)
            text_from_zip = ""
            try:
                import zipfile
                import xml.etree.ElementTree as ET
                
                with zipfile.ZipFile(docx_path, 'r') as docx:
                    # Lire le document principal
                    xml_content = docx.read('word/document.xml')
                    
                    # Parser le XML
                    root = ET.fromstring(xml_content)
                    
                    # Namespace pour Office Open XML
                    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    
                    # Extraire tout le texte
                    text_parts = []
                    for elem in root.findall('.//w:t', ns):
                        if elem.text:
                            text_parts.append(elem.text)
                    
                    text_from_zip = ' '.join(text_parts)
                    logger.info(f"✅ zipfile XML: {len(text_from_zip)} caractères")
                    
            except Exception as e:
                logger.error(f"❌ Erreur zipfile: {e}")
            
            # Choisir la meilleure extraction
            if text_from_zip and len(text_from_zip) > len(text_from_docx):
                text = text_from_zip
                confidence = 0.9  # Haute confiance pour zipfile
                logger.info(f"📊 Utilisation extraction zipfile ({len(text)} caractères)")
            elif text_from_docx:
                text = text_from_docx
                confidence = 1.0  # Confiance maximale pour python-docx
                logger.info(f"📊 Utilisation extraction python-docx ({len(text)} caractères)")
            else:
                text = ""
                confidence = 0.0
                logger.error("❌ Aucune méthode n'a extrait de texte")
            
            processing_time = time.time() - start_time
            
            if text:
                logger.info(f"✅ Extraction Word réussie: {len(text)} caractères en {processing_time:.3f}s")
                
                # Nettoyer le texte (enlever les espaces multiples)
                text = re.sub(r'\s+', ' ', text).strip()
                
                return text, confidence, processing_time
            else:
                return "", 0.0, processing_time
                
        except Exception as e:
            logger.error(f"❌ DOCX extraction failed: {e}", exc_info=True)
            return "", 0.0, time.time() - start_time
    
    def extract_from_excel(self, excel_path: str) -> Tuple[str, float, float]:
        """Extract text from Excel file"""
        try:
            import pandas as pd
            
            if excel_path.endswith('.xlsx'):
                df = pd.read_excel(excel_path, sheet_name=None, header=None)
            else:
                df = pd.read_excel(excel_path, sheet_name=None, header=None, engine='xlrd')
            
            text_parts = []
            for sheet_name, sheet_data in df.items():
                text_parts.append(f"Sheet: {sheet_name}")
                text_parts.append(sheet_data.to_string(index=False, header=False))
            
            return "\n".join(text_parts), 1.0, 0.0
            
        except Exception as e:
            logger.error(f"Excel extraction failed: {e}")
            return "", 0.0, 0.0
    
    def detect_file_type(self, file_path: str) -> str:
        """Detect file type using magic - VERSION CORRIGÉE"""
        try:
            mime = magic.Magic(mime=True)
            mime_type = mime.from_file(file_path)
            
            # AJOUTER CE LOG POUR DEBUG
            logger.info(f"🔍 Détection type: {file_path} -> {mime_type}")
            
            if 'image/' in mime_type:
                return 'image'
            elif 'pdf' in mime_type:
                return 'pdf'
            elif 'wordprocessingml' in mime_type or 'msword' in mime_type:
                return 'docx'
            elif 'spreadsheetml' in mime_type or 'excel' in mime_type:
                return 'excel'
            elif 'text/' in mime_type:
                return 'text'
            else:
                # Fallback: vérifier l'extension
                ext = os.path.splitext(file_path)[1].lower()
                if ext == '.docx':
                    return 'docx'
                elif ext == '.doc':
                    return 'docx'
                elif ext == '.pdf':
                    return 'pdf'
                elif ext in ['.jpg', '.jpeg', '.png', '.bmp', '.gif']:
                    return 'image'
                elif ext in ['.xlsx', '.xls']:
                    return 'excel'
                else:
                    return 'unknown'
                    
        except Exception as e:
            logger.error(f"File type detection failed: {e}")
            
            # FALLBACK IMPORTANT: Détection par extension
            logger.info(f"🔄 Fallback: détection par extension pour {file_path}")
            
            ext = os.path.splitext(file_path)[1].lower()
            if ext == '.docx' or ext == '.doc':
                return 'docx'
            elif ext == '.pdf':
                return 'pdf'
            elif ext in ['.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff', '.webp']:
                return 'image'
            elif ext in ['.xlsx', '.xls']:
                return 'excel'
            elif ext in ['.txt', '.csv', '.json', '.xml']:
                return 'text'
            else:
                return 'unknown'
    
    def process_document(self, file_path: str, language: str = None) -> Dict[str, Any]:
        """Main method to process any document type"""
        # AJOUTER CE LOG
        logger.info(f"🛠️ DEBUG - Fichier reçu: {file_path}")
        logger.info(f"🛠️ DEBUG - Fichier existe: {os.path.exists(file_path)}")
        logger.info(f"🛠️ DEBUG - Taille: {os.path.getsize(file_path) if os.path.exists(file_path) else 'N/A'} bytes")
        
        file_type = self.detect_file_type(file_path)
        start_time = time.time()
        
        # AJOUTER CE LOG
        logger.info(f"🛠️ DEBUG - Type détecté: {file_type}")
        
        try:
            if file_type == 'image':
                text, confidence, _ = self.extract_from_image(file_path, language)
                pages = [(text, confidence)]
            elif file_type == 'pdf':
                # AJOUTER CE LOG
                logger.info(f"🔍 DEBUG - Début extraction PDF avec PyMuPDF...")
                pages = self.extract_from_pdf(file_path, language)
                text = "\n".join([page[0] for page in pages])
                confidence = np.mean([page[1] for page in pages]) if pages else 0.0
                
                # AJOUTER CE LOG
                logger.info(f"🔍 DEBUG - Extraction PDF terminée. Pages: {len(pages)}, Texte: {len(text)} caractères")
            elif file_type == 'docx':
                text, confidence, _ = self.extract_from_docx(file_path)
                pages = [(text, confidence)]
            elif file_type == 'excel':
                text, confidence, _ = self.extract_from_excel(file_path)
                pages = [(text, confidence)]
            else:
                logger.error(f"❌ Type non supporté: {file_type}")
                raise ValueError(f"Type de fichier non supporté: {file_type}")
            
            processing_time = time.time() - start_time
            
            return {
                'success': True,
                'file_type': file_type,
                'text': text,
                'confidence': float(confidence),
                'pages': pages,
                'processing_time': processing_time,
                'error': None
            }
            
        except Exception as e:
            processing_time = time.time() - start_time
            logger.error(f"❌ Traitement document échoué: {e}")
            import traceback
            logger.error(f"❌ Traceback: {traceback.format_exc()}")
            
            return {
                'success': False,
                'file_type': file_type,
                'text': "",
                'confidence': 0.0,
                'pages': [],
                'processing_time': processing_time,
                'error': str(e)
            }